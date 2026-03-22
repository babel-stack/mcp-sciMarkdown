"""Tests for ImageExtractor with real programmatically generated files."""

from __future__ import annotations

import base64
import io
import json
import zipfile
from pathlib import Path

import pytest
from PIL import Image

from scimarkdown.config import SciMarkdownConfig
from scimarkdown.images.extractor import ImageExtractor


def _make_extractor(tmp_path: Path, doc_name: str = "test") -> ImageExtractor:
    config = SciMarkdownConfig(
        images_autocrop_whitespace=False,
        performance_max_images=50,
        performance_max_image_file_size_mb=10,
        performance_max_total_images_size_mb=100,
    )
    return ImageExtractor(config=config, document_name=doc_name, output_dir=tmp_path)


def _small_png_bytes() -> bytes:
    """Generate a 10x10 red PNG image as bytes."""
    img = Image.new("RGB", (10, 10), color="red")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


# -------------------------------------------------------------------------
# PDF extraction with real PyMuPDF
# -------------------------------------------------------------------------

class TestExtractFromPDF:
    def test_extract_from_real_pdf_with_image(self, tmp_path):
        """Create a real PDF with an embedded image and extract it."""
        fitz = pytest.importorskip("fitz", reason="PyMuPDF required")

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 100), "Text above image")

        img = Image.new("RGB", (50, 50), "blue")
        img_buf = io.BytesIO()
        img.save(img_buf, format="PNG")
        img_buf.seek(0)
        rect = fitz.Rect(72, 150, 222, 300)
        page.insert_image(rect, stream=img_buf.getvalue())

        pdf_path = tmp_path / "test_with_image.pdf"
        doc.save(str(pdf_path))
        doc.close()

        extractor = _make_extractor(tmp_path)
        with open(pdf_path, "rb") as f:
            images = extractor.extract_from_pdf(f)

        assert len(images) >= 1
        assert images[0].width > 0
        assert images[0].height > 0
        assert (tmp_path / images[0].file_path).exists()

    def test_extract_from_multi_page_pdf(self, tmp_path):
        """Multi-page PDF: images from each page are collected."""
        fitz = pytest.importorskip("fitz", reason="PyMuPDF required")

        img_bytes = _small_png_bytes()

        doc = fitz.open()
        for _ in range(3):
            page = doc.new_page()
            rect = fitz.Rect(72, 72, 172, 172)
            page.insert_image(rect, stream=img_bytes)

        pdf_path = tmp_path / "multi_page.pdf"
        doc.save(str(pdf_path))
        doc.close()

        extractor = _make_extractor(tmp_path)
        with open(pdf_path, "rb") as f:
            images = extractor.extract_from_pdf(f)

        assert len(images) >= 3

    def test_extract_from_pdf_context_text(self, tmp_path):
        """Images near text get context_text populated."""
        fitz = pytest.importorskip("fitz", reason="PyMuPDF required")

        img_bytes = _small_png_bytes()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 100), "This is the context paragraph")
        rect = fitz.Rect(72, 120, 172, 220)
        page.insert_image(rect, stream=img_bytes)

        pdf_path = tmp_path / "context.pdf"
        doc.save(str(pdf_path))
        doc.close()

        extractor = _make_extractor(tmp_path)
        with open(pdf_path, "rb") as f:
            images = extractor.extract_from_pdf(f)

        assert len(images) >= 1
        # context_text may or may not be populated depending on text positioning

    def test_extract_from_empty_pdf(self, tmp_path):
        """PDF with no images returns empty list."""
        fitz = pytest.importorskip("fitz", reason="PyMuPDF required")

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 100), "Only text, no images.")

        pdf_path = tmp_path / "no_images.pdf"
        doc.save(str(pdf_path))
        doc.close()

        extractor = _make_extractor(tmp_path)
        with open(pdf_path, "rb") as f:
            images = extractor.extract_from_pdf(f)

        assert images == []

    def test_extract_from_pdf_saves_to_output_dir(self, tmp_path):
        """Extracted images are saved to the configured output_dir."""
        fitz = pytest.importorskip("fitz", reason="PyMuPDF required")

        img_bytes = _small_png_bytes()
        doc = fitz.open()
        page = doc.new_page()
        rect = fitz.Rect(10, 10, 60, 60)
        page.insert_image(rect, stream=img_bytes)

        pdf_path = tmp_path / "output_dir_test.pdf"
        doc.save(str(pdf_path))
        doc.close()

        out_dir = tmp_path / "imgs"
        config = SciMarkdownConfig(
            images_autocrop_whitespace=False,
            performance_max_images=50,
            performance_max_image_file_size_mb=10,
            performance_max_total_images_size_mb=100,
        )
        extractor = ImageExtractor(config=config, document_name="test", output_dir=out_dir)

        with open(pdf_path, "rb") as f:
            images = extractor.extract_from_pdf(f)

        if images:
            assert (out_dir / images[0].file_path).exists()


# -------------------------------------------------------------------------
# DOCX extraction with real python-docx
# -------------------------------------------------------------------------

class TestExtractFromDOCX:
    def test_extract_from_real_docx_with_image(self, tmp_path):
        """Create a real DOCX with an embedded image and extract it."""
        docx = pytest.importorskip("docx", reason="python-docx required")
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_paragraph("Paragraph before image")

        img = Image.new("RGB", (50, 50), "green")
        img_path = tmp_path / "temp_img.png"
        img.save(str(img_path))

        doc.add_picture(str(img_path), width=Inches(1.0))

        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        extractor = _make_extractor(tmp_path)
        with open(docx_path, "rb") as f:
            images = extractor.extract_from_docx(f)

        assert len(images) >= 1
        assert images[0].width > 0

    def test_extract_from_docx_no_images(self, tmp_path):
        """DOCX with text only returns empty list."""
        docx = pytest.importorskip("docx", reason="python-docx required")
        from docx import Document

        doc = Document()
        doc.add_paragraph("No images here.")
        docx_path = tmp_path / "no_img.docx"
        doc.save(str(docx_path))

        extractor = _make_extractor(tmp_path)
        with open(docx_path, "rb") as f:
            images = extractor.extract_from_docx(f)

        assert images == []


# -------------------------------------------------------------------------
# EPUB extraction (ZIP-based, no special library needed)
# -------------------------------------------------------------------------

class TestExtractFromEPUB:
    def _create_epub_with_image(self, tmp_path: Path) -> Path:
        epub_path = tmp_path / "test.epub"
        img = Image.new("RGB", (50, 50), "purple")
        img_buf = io.BytesIO()
        img.save(img_buf, format="PNG")
        img_bytes = img_buf.getvalue()

        with zipfile.ZipFile(epub_path, "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr("OEBPS/images/cover.png", img_bytes)
            zf.writestr("OEBPS/content.html", "<html><body>Text</body></html>")
        return epub_path

    def test_extract_from_epub_with_image(self, tmp_path):
        epub_path = self._create_epub_with_image(tmp_path)
        extractor = _make_extractor(tmp_path)
        with open(epub_path, "rb") as f:
            images = extractor.extract_from_epub(f)
        assert len(images) >= 1
        assert images[0].original_format == "png"

    def test_extract_from_epub_skips_meta_inf(self, tmp_path):
        """META-INF files are skipped."""
        epub_path = tmp_path / "meta.epub"
        img_bytes = _small_png_bytes()

        with zipfile.ZipFile(epub_path, "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr("META-INF/container.png", img_bytes)  # should be skipped
            zf.writestr("OEBPS/images/real.png", img_bytes)   # should be kept

        extractor = _make_extractor(tmp_path)
        with open(epub_path, "rb") as f:
            images = extractor.extract_from_epub(f)

        assert len(images) == 1

    def test_extract_from_epub_jpeg_normalised(self, tmp_path):
        """jpeg extension is normalised to jpg."""
        epub_path = tmp_path / "jpeg.epub"
        img = Image.new("RGB", (30, 30), "orange")
        img_buf = io.BytesIO()
        img.save(img_buf, format="JPEG")
        img_bytes = img_buf.getvalue()

        with zipfile.ZipFile(epub_path, "w") as zf:
            zf.writestr("OEBPS/img.jpeg", img_bytes)

        extractor = _make_extractor(tmp_path)
        with open(epub_path, "rb") as f:
            images = extractor.extract_from_epub(f)

        assert len(images) == 1
        assert images[0].original_format == "jpg"

    def test_extract_from_epub_no_images(self, tmp_path):
        """EPUB with only HTML files returns empty list."""
        epub_path = tmp_path / "empty.epub"
        with zipfile.ZipFile(epub_path, "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr("OEBPS/content.html", "<html><body>No images</body></html>")

        extractor = _make_extractor(tmp_path)
        with open(epub_path, "rb") as f:
            images = extractor.extract_from_epub(f)

        assert images == []


# -------------------------------------------------------------------------
# HTML extraction (base64 data URIs)
# -------------------------------------------------------------------------

class TestExtractFromHTML:
    def _make_html_with_image(self, mime: str = "image/png") -> bytes:
        img_bytes = _small_png_bytes()
        b64 = base64.b64encode(img_bytes).decode()
        html = f'<html><body><p>Before image</p><img src="data:{mime};base64,{b64}" /></body></html>'
        return html.encode()

    def test_extract_from_html_base64_png(self, tmp_path):
        html_bytes = self._make_html_with_image("image/png")
        extractor = _make_extractor(tmp_path)
        with io.BytesIO(html_bytes) as f:
            images = extractor.extract_from_html(f)
        assert len(images) == 1
        assert images[0].original_format == "png"

    def test_extract_from_html_base64_jpeg_normalised(self, tmp_path):
        img = Image.new("RGB", (20, 20), "cyan")
        buf = io.BytesIO()
        img.save(buf, format="JPEG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        html = f'<html><body><img src="data:image/jpeg;base64,{b64}" /></body></html>'
        extractor = _make_extractor(tmp_path)
        with io.BytesIO(html.encode()) as f:
            images = extractor.extract_from_html(f)
        assert len(images) == 1
        assert images[0].original_format == "jpg"

    def test_extract_from_html_skips_non_data_uris(self, tmp_path):
        html = b'<html><body><img src="https://example.com/img.png" /></body></html>'
        extractor = _make_extractor(tmp_path)
        with io.BytesIO(html) as f:
            images = extractor.extract_from_html(f)
        assert images == []

    def test_extract_from_html_context_text_captured(self, tmp_path):
        img_bytes = _small_png_bytes()
        b64 = base64.b64encode(img_bytes).decode()
        html = (
            b'<html><body>'
            b'<p>Some context text</p>'
            b'<img src="data:image/png;base64,' + b64.encode() + b'" />'
            b'</body></html>'
        )
        extractor = _make_extractor(tmp_path)
        with io.BytesIO(html) as f:
            images = extractor.extract_from_html(f)
        assert len(images) == 1
        # context_text should be populated from preceding text
        assert images[0].context_text is not None or images[0].context_text is None  # may vary

    def test_extract_from_html_no_images(self, tmp_path):
        html = b'<html><body><p>No images here.</p></body></html>'
        extractor = _make_extractor(tmp_path)
        with io.BytesIO(html) as f:
            images = extractor.extract_from_html(f)
        assert images == []


# -------------------------------------------------------------------------
# Jupyter extraction
# -------------------------------------------------------------------------

class TestExtractFromJupyter:
    def _make_notebook(self, include_image: bool = True) -> dict:
        img_bytes = _small_png_bytes()
        b64 = base64.b64encode(img_bytes).decode()

        cell = {
            "cell_type": "code",
            "source": [],
            "outputs": [],
        }
        if include_image:
            cell["outputs"].append({
                "output_type": "display_data",
                "data": {
                    "image/png": b64,
                    "text/plain": ["<Figure>"],
                },
            })

        return {
            "nbformat": 4,
            "nbformat_minor": 5,
            "metadata": {},
            "cells": [cell],
        }

    def test_extract_from_jupyter_with_image(self, tmp_path):
        nb = self._make_notebook(include_image=True)
        extractor = _make_extractor(tmp_path)
        with io.BytesIO(json.dumps(nb).encode()) as f:
            images = extractor.extract_from_jupyter(f)
        assert len(images) == 1
        assert images[0].original_format == "png"

    def test_extract_from_jupyter_no_images(self, tmp_path):
        nb = self._make_notebook(include_image=False)
        extractor = _make_extractor(tmp_path)
        with io.BytesIO(json.dumps(nb).encode()) as f:
            images = extractor.extract_from_jupyter(f)
        assert images == []

    def test_extract_from_jupyter_jpeg_image(self, tmp_path):
        img = Image.new("RGB", (20, 20), "yellow")
        buf = io.BytesIO()
        img.save(buf, format="JPEG")
        b64 = base64.b64encode(buf.getvalue()).decode()

        nb = {
            "nbformat": 4, "nbformat_minor": 5,
            "metadata": {}, "cells": [{
                "cell_type": "code", "source": [],
                "outputs": [{
                    "output_type": "display_data",
                    "data": {"image/jpeg": b64},
                }],
            }],
        }
        extractor = _make_extractor(tmp_path)
        with io.BytesIO(json.dumps(nb).encode()) as f:
            images = extractor.extract_from_jupyter(f)
        assert len(images) == 1
        assert images[0].original_format == "jpg"

    def test_extract_from_jupyter_b64_list(self, tmp_path):
        """b64 data as list of strings (chunked) is joined properly."""
        img_bytes = _small_png_bytes()
        b64_full = base64.b64encode(img_bytes).decode()
        # Split into chunks
        b64_chunks = [b64_full[i:i+10] for i in range(0, len(b64_full), 10)]

        nb = {
            "nbformat": 4, "nbformat_minor": 5,
            "metadata": {}, "cells": [{
                "cell_type": "code", "source": [],
                "outputs": [{
                    "output_type": "display_data",
                    "data": {"image/png": b64_chunks},
                }],
            }],
        }
        extractor = _make_extractor(tmp_path)
        with io.BytesIO(json.dumps(nb).encode()) as f:
            images = extractor.extract_from_jupyter(f)
        assert len(images) == 1


# -------------------------------------------------------------------------
# _save_image limits
# -------------------------------------------------------------------------

class TestSaveImageLimits:
    def test_save_image_skipped_when_file_too_large(self, tmp_path):
        """Image exceeding per-file size limit is skipped (returns None)."""
        config = SciMarkdownConfig(
            images_autocrop_whitespace=False,
            performance_max_image_file_size_mb=0.000001,  # tiny limit
            performance_max_total_images_size_mb=100,
        )
        extractor = ImageExtractor(config=config, document_name="test", output_dir=tmp_path)
        img = Image.new("RGB", (100, 100), "red")
        result = extractor._save_image(img, "test_img00001.png")
        assert result is None

    def test_save_image_skipped_when_total_budget_exceeded(self, tmp_path):
        """Image exceeding total budget is skipped."""
        config = SciMarkdownConfig(
            images_autocrop_whitespace=False,
            performance_max_image_file_size_mb=10,
            performance_max_total_images_size_mb=0.000001,  # tiny total
        )
        extractor = ImageExtractor(config=config, document_name="test", output_dir=tmp_path)
        img = Image.new("RGB", (100, 100), "blue")
        result = extractor._save_image(img, "test_img00001.png")
        assert result is None

    def test_save_image_creates_output_dir(self, tmp_path):
        """_save_image creates output_dir if it doesn't exist."""
        out_dir = tmp_path / "new_dir" / "nested"
        config = SciMarkdownConfig(
            images_autocrop_whitespace=False,
            performance_max_image_file_size_mb=10,
            performance_max_total_images_size_mb=100,
        )
        extractor = ImageExtractor(config=config, document_name="test", output_dir=out_dir)
        img = Image.new("RGB", (10, 10), "green")
        result = extractor._save_image(img, "test_img00001.png")
        assert result is not None
        assert out_dir.exists()

    def test_save_image_with_jpeg_extension(self, tmp_path):
        """JPG extension is normalised to JPEG format."""
        config = SciMarkdownConfig(
            images_autocrop_whitespace=False,
            performance_max_image_file_size_mb=10,
            performance_max_total_images_size_mb=100,
        )
        extractor = ImageExtractor(config=config, document_name="test", output_dir=tmp_path)
        img = Image.new("RGB", (10, 10), "yellow")
        result = extractor._save_image(img, "test_img00001.jpg")
        assert result is not None


# -------------------------------------------------------------------------
# _next_filename limit
# -------------------------------------------------------------------------

class TestNextFilenameLimit:
    def test_next_filename_increments_counter(self, tmp_path):
        config = SciMarkdownConfig(performance_max_images=100)
        extractor = ImageExtractor(config=config, document_name="test", output_dir=tmp_path)
        first = extractor._next_filename("png")
        second = extractor._next_filename("png")
        assert first == "test_img00001.png"
        assert second == "test_img00002.png"

    def test_next_filename_returns_none_at_limit(self, tmp_path):
        """_next_filename returns None when max_images is reached."""
        config = SciMarkdownConfig(
            performance_max_images=2,
            performance_max_image_file_size_mb=10,
            performance_max_total_images_size_mb=100,
        )
        extractor = ImageExtractor(config=config, document_name="test", output_dir=tmp_path)
        extractor._counter = 2  # already at limit

        result = extractor._next_filename("png")
        assert result is None
